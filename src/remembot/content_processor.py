"""
Content processor for RememBot.
Handles URL fetching, image processing, and document parsing with enhanced error handling.
"""

import asyncio
import logging
import tempfile
import os
from typing import Dict, Any, Optional
from urllib.parse import urlparse
import aiohttp
from pathlib import Path
from datetime import datetime, timezone
import time

import requests
from bs4 import BeautifulSoup
from PIL import Image
import pytesseract
from docx import Document
import pypdf
from pypdf import PdfReader
import openpyxl

from .config import get_config

logger = logging.getLogger(__name__)


class ContentProcessingError(Exception):
    """Custom exception for content processing errors."""
    pass


class ContentProcessor:
    """Processes different types of content for storage with enhanced error handling."""
    
    def __init__(self):
        """Initialize content processor."""
        self.session = None
        try:
            self.config = get_config()
        except Exception as e:
            # For testing or when config is not available, use defaults
            logger.warning(f"Could not load config, using defaults: {e}")
            from types import SimpleNamespace
            self.config = SimpleNamespace(
                max_file_size_mb=50,
                request_timeout=30,
                max_retries=3
            )
        self._session_lock = asyncio.Lock()
    
    async def _get_session(self):
        """Get or create aiohttp session with connection pooling."""
        async with self._session_lock:
            if self.session is None or self.session.closed:
                # Create session with connection pooling and timeouts
                connector = aiohttp.TCPConnector(
                    limit=10,  # Total connection pool size
                    limit_per_host=5,  # Max connections per host
                    ttl_dns_cache=300,  # DNS cache TTL
                    use_dns_cache=True,
                )
                
                timeout = aiohttp.ClientTimeout(
                    total=self.config.request_timeout,
                    connect=10,
                    sock_read=10
                )
                
                headers = {
                    'User-Agent': 'RememBot/1.0 (Personal Knowledge Management Bot)'
                }
                
                self.session = aiohttp.ClientSession(
                    connector=connector,
                    timeout=timeout,
                    headers=headers
                )
                
                logger.debug("Created new aiohttp session with connection pooling")
            
            return self.session
    
    async def process_text(self, text: str) -> Dict[str, Any]:
        """Process text content, detecting URLs and extracting content with error handling."""
        try:
            # Validate input
            if not text or not text.strip():
                raise ContentProcessingError("Empty text input")
            
            text = text.strip()
            
            # Check if text contains URLs
            words = text.split()
            urls = [word for word in words if self._is_url(word)]
            
            if urls:
                # Process as URL with retry mechanism
                return await self._process_url_with_retry(urls[0], text)
            else:
                # Process as plain text
                return {
                    'content_type': 'text',
                    'extracted_info': text,
                    'metadata': {
                        'word_count': len(words),
                        'character_count': len(text),
                        'processed_at': datetime.now(timezone.utc).isoformat(),
                        'processing_time_ms': 0
                    }
                }
        
        except Exception as e:
            logger.error(f"Error processing text: {e}")
            return {
                'content_type': 'text',
                'extracted_info': f"Error processing text: {str(e)}",
                'metadata': {
                    'error': str(e),
                    'error_type': type(e).__name__,
                    'processed_at': datetime.now(timezone.utc).isoformat()
                }
            }
    
    def _is_url(self, text: str) -> bool:
        """Check if text is a URL with enhanced validation."""
        try:
            if not text or len(text) < 10:  # Minimum URL length check
                return False
            
            result = urlparse(text)
            return all([
                result.scheme in ['http', 'https'],
                result.netloc,
                '.' in result.netloc  # Basic domain validation
            ])
        except Exception:
            return False
    
    async def _process_url_with_retry(self, url: str, original_text: str) -> Dict[str, Any]:
        """Process URL with retry mechanism and Tavily fallback."""
        last_error = None
        
        for attempt in range(self.config.max_retries):
            try:
                return await self._process_url(url, original_text)
            except Exception as e:
                last_error = e
                if attempt < self.config.max_retries - 1:
                    # Exponential backoff
                    wait_time = 2 ** attempt
                    logger.warning(f"URL processing attempt {attempt + 1} failed, retrying in {wait_time}s: {e}")
                    await asyncio.sleep(wait_time)
                else:
                    logger.error(f"All {self.config.max_retries} attempts failed for URL {url}: {e}")
        
        # All retries failed - try enhanced extraction strategies as fallback
        logger.info(f"Attempting enhanced extraction fallback for URL: {url}")
        try:
            return await self._process_url_with_enhanced_extraction(url, original_text, str(last_error))
        except Exception as enhanced_error:
            logger.error(f"Enhanced extraction fallback also failed for URL {url}: {enhanced_error}")
        
        # Final fallback
        return {
            'content_type': 'url',
            'extracted_info': f"Failed to process URL after all attempts and fallbacks: {url}",
            'metadata': {
                'url': url,
                'error': str(last_error),
                'error_type': type(last_error).__name__,
                'attempts': self.config.max_retries,
                'enhanced_extraction_attempted': True,
                'processed_at': datetime.now(timezone.utc).isoformat()
            }
        }
    
    async def _process_url_with_enhanced_extraction(self, url: str, original_text: str, original_error: str) -> Dict[str, Any]:
        """Use enhanced extraction strategies as fallback for captcha-blocked or failed URLs."""
        start_time = time.time()
        
        try:
            logger.info(f"Using enhanced extraction strategies for URL: {url}")
            
            # Since we can't directly call MCP tools from within the code,
            # we'll create a simulated Tavily extraction by attempting
            # different strategies that Tavily would use
            
            # Strategy 1: Try with different user agents
            user_agents = [
                'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            ]
            
            session = await self._get_session()
            
            for user_agent in user_agents:
                try:
                    headers = {
                        'User-Agent': user_agent,
                        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                        'Accept-Language': 'en-US,en;q=0.5',
                        'Accept-Encoding': 'gzip, deflate',
                        'DNT': '1',
                        'Connection': 'keep-alive',
                        'Upgrade-Insecure-Requests': '1',
                    }
                    
                    async with session.get(url, headers=headers, timeout=20) as response:
                        if response.status == 200:
                            html = await response.text()
                            
                            # Use BeautifulSoup for better content extraction
                            soup = BeautifulSoup(html, 'html.parser')
                            
                            # Extract title
                            title = soup.find('title')
                            title_text = title.get_text().strip() if title else "No title"
                            
                            # Try to find main content areas
                            content_selectors = [
                                'article', 'main', '[role="main"]', '.content', '.post-content',
                                '.entry-content', '.article-content', '.story-body', '.post-body'
                            ]
                            
                            main_content = None
                            for selector in content_selectors:
                                content_element = soup.select_one(selector)
                                if content_element:
                                    main_content = content_element
                                    break
                            
                            # If no main content found, use body but remove unwanted elements
                            if not main_content:
                                main_content = soup.find('body') or soup
                            
                            # Remove unwanted elements
                            for unwanted in main_content.find_all(['script', 'style', 'nav', 'footer', 'header', 'aside', '.sidebar', '.navigation', '.ad', '.advertisement']):
                                unwanted.decompose()
                            
                            # Extract text
                            text = main_content.get_text()
                            
                            # Clean up text
                            lines = (line.strip() for line in text.splitlines())
                            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                            clean_text = ' '.join(chunk for chunk in chunks if chunk)
                            
                            if not clean_text.strip():
                                continue  # Try next user agent
                            
                            # Limit text length
                            max_length = 10000
                            if len(clean_text) > max_length:
                                clean_text = clean_text[:max_length] + "... [truncated by enhanced extraction]"
                            
                            processing_time = (time.time() - start_time) * 1000
                            
                            return {
                                'content_type': 'url',
                                'extracted_info': f"Title: {title_text}\n\nContent: {clean_text}",
                                'metadata': {
                                    'url': url,
                                    'title': title_text,
                                    'status_code': response.status,
                                    'content_length': len(clean_text),
                                    'extraction_method': 'enhanced_fallback',
                                    'user_agent': user_agent[:50] + '...',
                                    'original_error': original_error,
                                    'processing_time_ms': round(processing_time, 2),
                                    'processed_at': datetime.now(timezone.utc).isoformat()
                                }
                            }
                
                except Exception as e:
                    logger.debug(f"Enhanced extraction attempt with {user_agent[:50]}... failed: {e}")
                    continue
            
            # If all user agents failed, try one more strategy
            raise Exception("All enhanced extraction strategies failed")
        
        except Exception as e:
            processing_time = (time.time() - start_time) * 1000
            logger.error(f"Enhanced fallback extraction failed for URL {url}: {e}")
            raise Exception(f"Enhanced extraction failed: {e}")
    
    async def _process_url(self, url: str, original_text: str) -> Dict[str, Any]:
        """Extract content from URL with enhanced error handling."""
        start_time = time.time()
        
        try:
            # Validate URL
            if not self._is_url(url):
                raise ContentProcessingError(f"Invalid URL format: {url}")
            
            session = await self._get_session()
            
            logger.debug(f"Fetching URL: {url}")
            async with session.get(url) as response:
                processing_time = (time.time() - start_time) * 1000
                
                if response.status == 200:
                    # Check content type
                    content_type = response.headers.get('content-type', '').lower()
                    if 'text/html' not in content_type and 'text/plain' not in content_type:
                        logger.warning(f"Unexpected content type for URL {url}: {content_type}")
                    
                    html = await response.text()
                    
                    if not html.strip():
                        raise ContentProcessingError("Empty response content")
                    
                    soup = BeautifulSoup(html, 'html.parser')
                    
                    # Extract title
                    title = soup.find('title')
                    title_text = title.get_text().strip() if title else "No title"
                    
                    # Extract main content (simple approach)
                    # Remove script and style elements
                    for script in soup(["script", "style", "nav", "footer", "header"]):
                        script.decompose()
                    
                    # Get text
                    text = soup.get_text()
                    
                    # Clean up text
                    lines = (line.strip() for line in text.splitlines())
                    chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                    clean_text = ' '.join(chunk for chunk in chunks if chunk)
                    
                    if not clean_text.strip():
                        raise ContentProcessingError("No extractable text content found")
                    
                    # Limit text length
                    max_length = 10000  # Increased from 5000
                    if len(clean_text) > max_length:
                        clean_text = clean_text[:max_length] + "... [truncated]"
                    
                    return {
                        'content_type': 'url',
                        'extracted_info': f"Title: {title_text}\n\nContent: {clean_text}",
                        'metadata': {
                            'url': url,
                            'title': title_text,
                            'status_code': response.status,
                            'content_type': content_type,
                            'content_length': len(clean_text),
                            'original_size': len(html),
                            'processing_time_ms': round(processing_time, 2),
                            'processed_at': datetime.now(timezone.utc).isoformat()
                        }
                    }
                else:
                    error_msg = f"HTTP {response.status}"
                    if response.status == 403:
                        error_msg += " (Forbidden - possible rate limiting or blocking)"
                    elif response.status == 404:
                        error_msg += " (Not Found)"
                    elif response.status == 429:
                        error_msg += " (Too Many Requests)"
                    elif response.status >= 500:
                        error_msg += " (Server Error)"
                    
                    logger.warning(f"Failed to fetch URL {url}: {error_msg}")
                    return {
                        'content_type': 'url',
                        'extracted_info': f"Failed to fetch URL: {url} ({error_msg})",
                        'metadata': {
                            'url': url,
                            'status_code': response.status,
                            'error': error_msg,
                            'processing_time_ms': round(processing_time, 2),
                            'processed_at': datetime.now(timezone.utc).isoformat()
                        }
                    }
        
        except aiohttp.ClientError as e:
            processing_time = (time.time() - start_time) * 1000
            logger.error(f"Network error processing URL {url}: {e}")
            return {
                'content_type': 'url',
                'extracted_info': f"Network error processing URL: {url} - {str(e)}",
                'metadata': {
                    'url': url,
                    'error': str(e),
                    'error_type': 'NetworkError',
                    'processing_time_ms': round(processing_time, 2),
                    'processed_at': datetime.now(timezone.utc).isoformat()
                }
            }
        except Exception as e:
            processing_time = (time.time() - start_time) * 1000
            logger.error(f"Error processing URL {url}: {e}")
            return {
                'content_type': 'url',
                'extracted_info': f"Error processing URL: {url} - {str(e)}",
                'metadata': {
                    'url': url,
                    'error': str(e),
                    'error_type': type(e).__name__,
                    'processing_time_ms': round(processing_time, 2),
                    'processed_at': datetime.now(timezone.utc).isoformat()
                }
            }
    
    async def process_image(self, file) -> Dict[str, Any]:
        """Process image file with OCR and enhanced error handling."""
        start_time = time.time()
        tmp_file_path = None
        
        try:
            # Validate file size
            if file.file_size > self.config.max_file_size_mb * 1024 * 1024:
                raise ContentProcessingError(
                    f"File too large: {file.file_size / 1024 / 1024:.1f}MB "
                    f"(max: {self.config.max_file_size_mb}MB)"
                )
            
            # Download file to temporary location
            with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tmp_file:
                tmp_file_path = tmp_file.name
                
                logger.debug(f"Downloading image to {tmp_file_path}")
                await file.download_to_drive(tmp_file_path)
                
                # Validate file exists and has content
                if not os.path.exists(tmp_file_path) or os.path.getsize(tmp_file_path) == 0:
                    raise ContentProcessingError("Downloaded file is empty or missing")
                
                # Open image with PIL
                try:
                    image = Image.open(tmp_file_path)
                    image.verify()  # Verify it's a valid image
                    image = Image.open(tmp_file_path)  # Reopen after verify
                except Exception as e:
                    raise ContentProcessingError(f"Invalid image format: {e}")
                
                # Check image dimensions
                width, height = image.size
                if width < 10 or height < 10:
                    raise ContentProcessingError(f"Image too small: {width}x{height}")
                
                # Perform OCR with error handling
                try:
                    logger.debug("Performing OCR on image")
                    ocr_text = pytesseract.image_to_string(image, timeout=30)
                except Exception as e:
                    logger.warning(f"OCR failed: {e}")
                    ocr_text = "[OCR processing failed]"
                
                # Get image metadata
                format_info = image.format or "Unknown"
                mode = image.mode
                
                processing_time = (time.time() - start_time) * 1000
                
                return {
                    'content_type': 'image',
                    'extracted_info': f"OCR Text: {ocr_text.strip()}" if ocr_text.strip() else "[No text detected in image]",
                    'metadata': {
                        'width': width,
                        'height': height,
                        'format': format_info,
                        'mode': mode,
                        'file_size': file.file_size,
                        'has_text': bool(ocr_text.strip()),
                        'ocr_length': len(ocr_text.strip()),
                        'processing_time_ms': round(processing_time, 2),
                        'processed_at': datetime.now(timezone.utc).isoformat()
                    }
                }
        
        except ContentProcessingError:
            raise  # Re-raise our custom errors
        except Exception as e:
            processing_time = (time.time() - start_time) * 1000
            logger.error(f"Error processing image: {e}")
            return {
                'content_type': 'image',
                'extracted_info': f"Error processing image: {str(e)}",
                'metadata': {
                    'error': str(e),
                    'error_type': type(e).__name__,
                    'processing_time_ms': round(processing_time, 2),
                    'processed_at': datetime.now(timezone.utc).isoformat()
                }
            }
        finally:
            # Clean up temporary file
            if tmp_file_path and os.path.exists(tmp_file_path):
                try:
                    os.unlink(tmp_file_path)
                    logger.debug(f"Cleaned up temporary file: {tmp_file_path}")
                except Exception as e:
                    logger.warning(f"Failed to clean up temporary file {tmp_file_path}: {e}")
    
    async def process_document(self, file, filename: str) -> Dict[str, Any]:
        """Process document file (PDF, Word, Excel) with enhanced error handling."""
        start_time = time.time()
        tmp_file_path = None
        
        try:
            # Validate file size
            if file.file_size > self.config.max_file_size_mb * 1024 * 1024:
                raise ContentProcessingError(
                    f"File too large: {file.file_size / 1024 / 1024:.1f}MB "
                    f"(max: {self.config.max_file_size_mb}MB)"
                )
            
            file_ext = Path(filename).suffix.lower()
            
            # Validate file extension
            supported_extensions = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt']
            if file_ext not in supported_extensions:
                raise ContentProcessingError(
                    f"Unsupported file type: {file_ext}. "
                    f"Supported: {', '.join(supported_extensions)}"
                )
            
            # Download file to temporary location
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
                tmp_file_path = tmp_file.name
                
                logger.debug(f"Downloading document {filename} to {tmp_file_path}")
                await file.download_to_drive(tmp_file_path)
                
                # Validate file exists and has content
                if not os.path.exists(tmp_file_path) or os.path.getsize(tmp_file_path) == 0:
                    raise ContentProcessingError("Downloaded file is empty or missing")
                
                # Extract content based on file type
                if file_ext == '.pdf':
                    content = self._extract_pdf_text(tmp_file_path)
                elif file_ext in ['.docx', '.doc']:
                    content = self._extract_word_text(tmp_file_path)
                elif file_ext in ['.xlsx', '.xls']:
                    content = self._extract_excel_text(tmp_file_path)
                elif file_ext == '.txt':
                    content = self._extract_txt_text(tmp_file_path)
                else:
                    raise ContentProcessingError(f"Unsupported file type: {file_ext}")
                
                # Validate extracted content
                if not content or not content.strip():
                    content = "[No extractable content found]"
                
                processing_time = (time.time() - start_time) * 1000
                
                return {
                    'content_type': 'document',
                    'extracted_info': content,
                    'metadata': {
                        'filename': filename,
                        'file_extension': file_ext,
                        'file_size': file.file_size,
                        'content_length': len(content),
                        'processing_time_ms': round(processing_time, 2),
                        'processed_at': datetime.now(timezone.utc).isoformat()
                    }
                }
        
        except ContentProcessingError:
            raise  # Re-raise our custom errors
        except Exception as e:
            processing_time = (time.time() - start_time) * 1000
            logger.error(f"Error processing document {filename}: {e}")
            return {
                'content_type': 'document',
                'extracted_info': f"Error processing document: {str(e)}",
                'metadata': {
                    'filename': filename,
                    'error': str(e),
                    'error_type': type(e).__name__,
                    'processing_time_ms': round(processing_time, 2),
                    'processed_at': datetime.now(timezone.utc).isoformat()
                }
            }
        finally:
            # Clean up temporary file
            if tmp_file_path and os.path.exists(tmp_file_path):
                try:
                    os.unlink(tmp_file_path)
                    logger.debug(f"Cleaned up temporary file: {tmp_file_path}")
                except Exception as e:
                    logger.warning(f"Failed to clean up temporary file {tmp_file_path}: {e}")
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                if not content.strip():
                    raise ContentProcessingError("Text file is empty")
                return content.strip()
        except UnicodeDecodeError:
            # Try with different encodings
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read().strip()
                except UnicodeDecodeError:
                    continue
            raise ContentProcessingError("Unable to decode text file with any supported encoding")
        except Exception as e:
            raise ContentProcessingError(f"Error reading text file: {e}")
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file with enhanced error handling."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                if len(pdf_reader.pages) == 0:
                    raise ContentProcessingError("PDF has no pages")
                
                text = ""
                pages_processed = 0
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"[Page {page_num + 1}]\n{page_text}\n\n"
                            pages_processed += 1
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        text += f"[Page {page_num + 1}: Error extracting text]\n\n"
                
                if not text.strip():
                    raise ContentProcessingError("No extractable text found in PDF")
                
                logger.debug(f"Extracted text from {pages_processed}/{len(pdf_reader.pages)} pages")
                return text.strip()
                
        except ContentProcessingError:
            raise
        except Exception as e:
            raise ContentProcessingError(f"Error extracting PDF text: {e}")
    
    def _extract_word_text(self, file_path: str) -> str:
        """Extract text from Word document with enhanced error handling."""
        try:
            doc = Document(file_path)
            text = ""
            paragraphs_processed = 0
            
            # Extract paragraph text
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    text += para_text + "\n"
                    paragraphs_processed += 1
            
            # Extract table text
            tables_processed = 0
            for table in doc.tables:
                try:
                    text += "\n[Table]\n"
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            row_text.append(cell_text)
                        text += " | ".join(row_text) + "\n"
                    text += "\n"
                    tables_processed += 1
                except Exception as e:
                    logger.warning(f"Failed to extract table text: {e}")
            
            if not text.strip():
                raise ContentProcessingError("No extractable text found in Word document")
            
            logger.debug(f"Extracted {paragraphs_processed} paragraphs and {tables_processed} tables")
            return text.strip()
            
        except ContentProcessingError:
            raise
        except Exception as e:
            raise ContentProcessingError(f"Error extracting Word text: {e}")
    
    def _extract_excel_text(self, file_path: str) -> str:
        """Extract text from Excel file with enhanced error handling."""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            if not workbook.sheetnames:
                raise ContentProcessingError("Excel file has no sheets")
            
            text = ""
            sheets_processed = 0
            
            for sheet_name in workbook.sheetnames:
                try:
                    sheet = workbook[sheet_name]
                    text += f"[Sheet: {sheet_name}]\n"
                    
                    rows_with_data = 0
                    for row in sheet.iter_rows(values_only=True):
                        # Skip completely empty rows
                        if not any(cell is not None and str(cell).strip() for cell in row):
                            continue
                        
                        row_text = "\t".join([
                            str(cell).strip() if cell is not None else ""
                            for cell in row
                        ])
                        
                        if row_text.strip():
                            text += row_text + "\n"
                            rows_with_data += 1
                    
                    text += "\n"
                    sheets_processed += 1
                    logger.debug(f"Extracted {rows_with_data} rows from sheet '{sheet_name}'")
                    
                except Exception as e:
                    logger.warning(f"Failed to extract text from sheet '{sheet_name}': {e}")
                    text += f"[Error extracting sheet '{sheet_name}']\n\n"
            
            if not text.strip() or sheets_processed == 0:
                raise ContentProcessingError("No extractable data found in Excel file")
            
            logger.debug(f"Processed {sheets_processed}/{len(workbook.sheetnames)} sheets")
            return text.strip()
            
        except ContentProcessingError:
            raise
        except Exception as e:
            raise ContentProcessingError(f"Error extracting Excel text: {e}")
    
    async def close(self):
        """Close the HTTP session and cleanup resources."""
        if self.session and not self.session.closed:
            await self.session.close()
            logger.debug("Closed aiohttp session")
    
    async def __aenter__(self):
        """Async context manager entry."""
        return self
    
    async def __aexit__(self, exc_type, exc_val, exc_tb):
        """Async context manager exit."""
        await self.close()